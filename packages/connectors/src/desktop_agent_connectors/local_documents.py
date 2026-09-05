"""Safe, bounded loading of supported local text documents."""

import hashlib
import io
from pathlib import Path

from desktop_agent_rag_core import Document
from docx import Document as DocxDocument
from pypdf import PdfReader


class LocalDocumentLoader:
    """Read text-like files strictly beneath one configured directory."""

    supported_suffixes = frozenset({".csv", ".docx", ".json", ".md", ".pdf", ".rst", ".txt"})

    def __init__(self, root: Path, *, max_bytes: int) -> None:
        self.root = root.resolve()
        self.max_bytes = max_bytes

    def scan(self) -> list[Document]:
        documents: list[Document] = []
        if not self.root.is_dir():
            return documents
        for path in sorted(self.root.rglob("*")):
            if (
                path.is_symlink()
                or not path.is_file()
                or path.suffix.lower() not in self.supported_suffixes
            ):
                continue
            resolved = path.resolve()
            if not resolved.is_relative_to(self.root) or resolved.stat().st_size > self.max_bytes:
                continue
            raw = resolved.read_bytes()
            checksum = hashlib.sha256(raw).hexdigest()
            relative = resolved.relative_to(self.root).as_posix()
            try:
                text = self._extract(raw, resolved.suffix.lower())
            except (OSError, UnicodeError, ValueError):
                continue
            documents.append(Document(checksum, relative, text, checksum))
        return documents

    @staticmethod
    def _extract(raw: bytes, suffix: str) -> str:
        if suffix == ".pdf":
            reader = PdfReader(io.BytesIO(raw))
            if reader.is_encrypted:
                raise ValueError("encrypted PDFs are not supported")
            return "\n\n".join(page.extract_text() or "" for page in reader.pages)
        if suffix == ".docx":
            document = DocxDocument(io.BytesIO(raw))
            return "\n".join(paragraph.text for paragraph in document.paragraphs)
        return raw.decode("utf-8")
